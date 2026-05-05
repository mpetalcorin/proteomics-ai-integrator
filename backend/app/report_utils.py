import os
import tempfile
import base64
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches


def save_base64_png(base64_string):
    if not base64_string:
        return None

    image_bytes = base64.b64decode(base64_string)

    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    temp.write(image_bytes)
    temp.close()

    return temp.name


def generate_pdf_report(advanced_result):
    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    temp.close()

    doc = SimpleDocTemplate(temp.name, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Proteomics AI Integrator Report", styles["Title"]))
    story.append(Spacer(1, 12))

    story.append(
        Paragraph(
            f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            styles["Normal"],
        )
    )

    story.append(Spacer(1, 12))

    story.append(Paragraph("Overview", styles["Heading2"]))
    story.append(
        Paragraph(
            f"This report analysed {advanced_result.get('n_features', 0)} molecular features. "
            f"Detected control samples: {', '.join(advanced_result.get('control_samples', []))}. "
            f"Detected case samples: {', '.join(advanced_result.get('case_samples', []))}.",
            styles["Normal"],
        )
    )

    story.append(Spacer(1, 12))

    volcano_path = save_base64_png(advanced_result.get("volcano_png_base64"))

    if volcano_path:
        story.append(Paragraph("Volcano Plot", styles["Heading2"]))
        story.append(Image(volcano_path, width=450, height=340))
        story.append(Spacer(1, 12))

    story.append(Paragraph("Top Biomarkers", styles["Heading2"]))

    rows = [["Feature", "log2FC", "p-value", "Score"]]

    for item in advanced_result.get("top_biomarkers", [])[:15]:
        rows.append(
            [
                str(item.get("feature", ""))[:30],
                f"{item.get('log2fc', 0):.3f}",
                f"{item.get('p_value', 1):.3e}",
                f"{item.get('biomarker_score', 0):.3f}",
            ]
        )

    table = Table(rows, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
            ]
        )
    )

    story.append(table)
    story.append(Spacer(1, 12))

    story.append(Paragraph("Pathway Enrichment", styles["Heading2"]))

    pathways = advanced_result.get("pathway_enrichment", [])

    if pathways:
        for pathway in pathways:
            story.append(
                Paragraph(
                    f"<b>{pathway.get('pathway')}</b>: "
                    f"{', '.join(pathway.get('matched_features', []))}",
                    styles["Normal"],
                )
            )
            story.append(Spacer(1, 6))
    else:
        story.append(
            Paragraph(
                "No pathway matches were detected using the lightweight built-in pathway engine.",
                styles["Normal"],
            )
        )

    story.append(Spacer(1, 12))
    story.append(Paragraph("Explainable AI", styles["Heading2"]))

    explainability = advanced_result.get("explainability", {})

    if "warning" in explainability:
        story.append(Paragraph(explainability["warning"], styles["Normal"]))
    else:
        story.append(
            Paragraph(
                f"Estimated model accuracy: "
                f"{explainability.get('model_accuracy_mean', 0):.2%}.",
                styles["Normal"],
            )
        )

        for item in explainability.get("top_explainable_features", [])[:10]:
            story.append(
                Paragraph(
                    f"{item.get('feature')}: importance "
                    f"{item.get('importance_mean', 0):.4f}",
                    styles["Normal"],
                )
            )

    doc.build(story)

    if volcano_path and os.path.exists(volcano_path):
        os.remove(volcano_path)

    return temp.name


def generate_word_report(advanced_result):
    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    temp.close()

    document = Document()

    document.add_heading("Proteomics AI Integrator Report", 0)
    document.add_paragraph(
        f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )

    document.add_heading("Overview", level=1)
    document.add_paragraph(
        f"This report analysed {advanced_result.get('n_features', 0)} molecular features."
    )
    document.add_paragraph(
        f"Control samples: {', '.join(advanced_result.get('control_samples', []))}"
    )
    document.add_paragraph(
        f"Case samples: {', '.join(advanced_result.get('case_samples', []))}"
    )

    volcano_path = save_base64_png(advanced_result.get("volcano_png_base64"))

    if volcano_path:
        document.add_heading("Volcano Plot", level=1)
        document.add_picture(volcano_path, width=Inches(6.0))

    document.add_heading("Top Biomarkers", level=1)

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = "Feature"
    header_cells[1].text = "log2FC"
    header_cells[2].text = "p-value"
    header_cells[3].text = "Score"

    for item in advanced_result.get("top_biomarkers", [])[:25]:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.get("feature", ""))[:40]
        row_cells[1].text = f"{item.get('log2fc', 0):.3f}"
        row_cells[2].text = f"{item.get('p_value', 1):.3e}"
        row_cells[3].text = f"{item.get('biomarker_score', 0):.3f}"

    document.add_heading("Pathway Enrichment", level=1)

    pathways = advanced_result.get("pathway_enrichment", [])

    if pathways:
        for pathway in pathways:
            document.add_paragraph(
                f"{pathway.get('pathway')}: "
                f"{', '.join(pathway.get('matched_features', []))}"
            )
    else:
        document.add_paragraph(
            "No pathway matches were detected using the lightweight built-in pathway engine."
        )

    document.add_heading("Explainable AI", level=1)

    explainability = advanced_result.get("explainability", {})

    if "warning" in explainability:
        document.add_paragraph(explainability["warning"])
    else:
        document.add_paragraph(
            f"Estimated model accuracy: "
            f"{explainability.get('model_accuracy_mean', 0):.2%}."
        )

        for item in explainability.get("top_explainable_features", [])[:15]:
            document.add_paragraph(
                f"{item.get('feature')}: importance "
                f"{item.get('importance_mean', 0):.4f}"
            )

    document.add_heading("Important Interpretation Note", level=1)
    document.add_paragraph(
        "This report is intended for exploratory analysis. Biomarkers and pathways "
        "should be validated using independent datasets, orthogonal assays, and "
        "appropriate biological controls."
    )

    document.save(temp.name)

    if volcano_path and os.path.exists(volcano_path):
        os.remove(volcano_path)

    return temp.name